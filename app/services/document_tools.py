"""
Document parsing tools — give the agent first-class read access to
PDF / Word / Excel and a multimodal "look at this page/image" fallback.

Two tools live here:

1. ``read_document(path)`` — structured text extraction.

   - ``.pdf``  → pypdfium2 全文文本，每页前注 ``[Page N]``
   - ``.docx`` → python-docx，按段落保留标题层级（``# / ## / ###``）
   - ``.xlsx`` → openpyxl 转每个 sheet 为 markdown 表格
   - 纯文本格式 (``.txt`` / ``.md`` / ``.csv`` / ``.json`` / ``.yaml`` / ``.log``)
     → 提示用 ``read_file``（避免与内置工具重复）
   - 输出超过 ``_MAX_TEXT_BYTES`` 截断，提示再用 ``view_pdf_page_or_image``
     或写脚本切片精读。

2. ``view_pdf_page_or_image(path, page=1)`` — 多模态视觉弥补。

   - 图片（``.png/.jpg/.jpeg/.webp/.gif``）→ 直接回 base64 image_url block
   - PDF → pypdfium2 渲染 **指定单页** 为 PNG → base64
   - **运行时 throttle**：同一进程线程同一 path 调用次数超过
     ``_MAX_VIEW_PER_PATH`` (=5) 即拒绝，提示改用 ``read_document``。
   - 不限制全局总次数（多文件场景 agent 可以各自 5 次）。

设计要点：
- 两个工具都是主进程直接调用（不进 script_runner 沙箱），但 path 全部
  经 ``safe_join`` / ``_is_allowed`` 校验，禁止穿越用户根目录 / 越权访问。
- 视觉工具 return ``list[dict]`` —— LangChain 0.3+ ToolNode 会自动包成
  ``ToolMessage(content=list)``，OpenAI / Anthropic / Gemini 都吃这种
  multimodal content 数组（``image_url`` block）。
- throttle 用 ``contextvars.ContextVar``（线程隔离），每个新 chat thread
  里的 agent 调用相互独立；同一 thread 内多个 LangGraph 节点共享计数。
- 凡是 read 都按需缓冲：pypdfium2 对单 page 用 ``len(pdf)`` 早返，
  避免大 PDF 全量解析；python-docx / openpyxl 的对象本身就是惰性的。
"""

from __future__ import annotations

import base64
import contextvars
import io
import os
from typing import List, Optional

from langchain_core.tools import tool

from app.core.path_security import safe_join
from app.core.security import get_user_filesystem_dir


# ──────────────── 配置常量 ────────────────

# read_document 输出截断阈值。超过即截断，提示 agent 用 view_pdf_page_or_image
# 或写脚本分段读。200KB 在 GPT-4o / Claude Sonnet 上约 50K-70K tokens，
# 是绝大多数对话单 turn 的合理上限。
_MAX_TEXT_BYTES = 200 * 1024

# view_pdf_page_or_image 单 thread / 单 path 调用次数上限。
# Agent 想"翻很多页"应该改用 read_document（一次拿全文）；这个工具
# 是给"看图表"/"OCR 失败时弥补"用的。
_MAX_VIEW_PER_PATH = 5

# 单页/单图渲染分辨率上限（pypdfium2 scale 1.5 ≈ 108 dpi，
# 对中文 PDF 文字勉强可读、表格清晰）。再高 base64 体积爆炸。
_PDF_RENDER_SCALE = 1.5

# 图片读入后限制原始 bytes 大小，避免 agent 把 100MB tiff 塞进对话。
_MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 8MB raw

# 纯文本格式：让 agent 改用内置 read_file 工具，不重复造轮子。
_PLAIN_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".jsonl",
    ".yaml", ".yml", ".log", ".ini", ".cfg", ".toml", ".xml",
    ".html", ".htm", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".go", ".rs", ".java", ".c", ".cpp", ".h", ".sh", ".bash",
    ".sql", ".env",
}

_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


# ──────────────── throttle (per-chat-thread) ────────────────

# 用 contextvars 而不是 threading.local：LangGraph 在 asyncio 里跑，每个
# 请求一个 contextvars copy；线程池 worker 也会继承父 context。这样同一
# 用户 thread 的多次 tool 调用（无论同步异步）共享同一计数 dict，不同
# 用户 / 不同对话之间天然隔离（因为 chat.py 每次请求都新建 context）。
_view_call_counts: contextvars.ContextVar[Optional[dict]] = contextvars.ContextVar(
    "_doc_view_counts", default=None
)


def _bump_and_check_view_count(path: str) -> Optional[str]:
    """Increment the per-path view counter for current thread/context.

    Returns an error message string if over limit, else None.
    """
    counts = _view_call_counts.get()
    if counts is None:
        counts = {}
        _view_call_counts.set(counts)
    cur = counts.get(path, 0) + 1
    counts[path] = cur
    if cur > _MAX_VIEW_PER_PATH:
        return (
            f"已对 {path} 调用 view_pdf_page_or_image {cur} 次（上限 "
            f"{_MAX_VIEW_PER_PATH}）。请改用 read_document 一次性读取文本，"
            "或写脚本（run_script）切片处理整本 PDF。视觉工具仅用于："
            "(a) 文本提取失败的扫描件；(b) 需要查看图表/插图细节。"
        )
    return None


def reset_view_throttle() -> None:
    """Reset throttle state — call at start of each new chat turn if needed.

    Currently we do NOT call this anywhere automatically; the contextvars
    isolation between requests is enough. Exposed for tests.
    """
    _view_call_counts.set(None)


# ──────────────── path 校验 ────────────────

def _resolve_path(fs_dir: str, path: str) -> tuple[Optional[str], Optional[str]]:
    """Returns (resolved_abs_path, error_message). Exactly one is set."""
    if not path:
        return None, "path 不能为空"
    try:
        resolved = safe_join(fs_dir, path)
    except PermissionError:
        return None, f"路径越权: {path}"
    if not os.path.exists(resolved):
        return None, f"文件不存在: {path}"
    if not os.path.isfile(resolved):
        return None, f"不是文件: {path}"
    return resolved, None


# ──────────────── 各格式提取器 ────────────────

def _extract_pdf_text(abs_path: str, max_bytes: int = _MAX_TEXT_BYTES) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return "[错误] pypdfium2 未安装，无法解析 PDF（请联系管理员安装依赖）"

    try:
        pdf = pdfium.PdfDocument(abs_path)
    except Exception as exc:  # pypdfium2 抛 PdfiumError 等
        return f"[错误] 无法打开 PDF: {exc}"

    try:
        n_pages = len(pdf)
        chunks: list[str] = []
        total_bytes = 0
        for i in range(n_pages):
            try:
                page = pdf[i]
                tp = page.get_textpage()
                text = tp.get_text_range() or ""
                tp.close()
                page.close()
            except Exception as exc:
                text = f"[第 {i+1} 页解析失败: {exc}]"
            chunk = f"\n[Page {i+1}/{n_pages}]\n{text}\n"
            chunk_bytes = chunk.encode("utf-8")
            if total_bytes + len(chunk_bytes) > max_bytes:
                chunks.append(
                    f"\n[已截断：剩余 {n_pages - i} 页未读取，"
                    "请用 view_pdf_page_or_image(path, page=N) 精读单页，"
                    "或 run_script 写 pypdfium2 脚本分段处理]"
                )
                break
            chunks.append(chunk)
            total_bytes += len(chunk_bytes)
        return "".join(chunks).strip() or "[PDF 没有可提取的文本（可能是扫描件，请用 view_pdf_page_or_image 查看）]"
    finally:
        try:
            pdf.close()
        except Exception:
            pass


def _extract_docx_text(abs_path: str, max_bytes: int = _MAX_TEXT_BYTES) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "[错误] python-docx 未安装，无法解析 .docx"

    try:
        doc = docx.Document(abs_path)
    except Exception as exc:
        return f"[错误] 无法打开 .docx: {exc}"

    chunks: list[str] = []
    total_bytes = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        # 标题层级 → markdown #
        style_name = (getattr(para.style, "name", "") or "").lower()
        prefix = ""
        if style_name.startswith("heading 1") or style_name == "title":
            prefix = "# "
        elif style_name.startswith("heading 2"):
            prefix = "## "
        elif style_name.startswith("heading 3"):
            prefix = "### "
        elif style_name.startswith("heading "):
            prefix = "#### "
        line = f"{prefix}{text}\n"
        line_bytes = line.encode("utf-8")
        if total_bytes + line_bytes.__len__() > max_bytes:
            chunks.append("\n[已截断：剩余段落未读取]")
            break
        chunks.append(line)
        total_bytes += len(line_bytes)

    # 表格也提取（如果还有空间）
    for tbl in doc.tables:
        if total_bytes >= max_bytes:
            break
        chunks.append("\n")
        for row in tbl.rows:
            cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            line = "| " + " | ".join(cells) + " |\n"
            line_bytes = line.encode("utf-8")
            if total_bytes + len(line_bytes) > max_bytes:
                chunks.append("[已截断：表格剩余行未读取]\n")
                total_bytes = max_bytes
                break
            chunks.append(line)
            total_bytes += len(line_bytes)
        chunks.append("\n")

    return "".join(chunks).strip() or "[.docx 没有可提取的文本]"


def _extract_xlsx_text(abs_path: str, max_bytes: int = _MAX_TEXT_BYTES) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[错误] openpyxl 未安装"

    try:
        wb = load_workbook(abs_path, data_only=True, read_only=True)
    except Exception as exc:
        return f"[错误] 无法打开 .xlsx: {exc}"

    chunks: list[str] = []
    total_bytes = 0
    truncated = False
    for sheet_name in wb.sheetnames:
        if truncated:
            break
        ws = wb[sheet_name]
        header = f"\n## Sheet: {sheet_name}\n\n"
        chunks.append(header)
        total_bytes += len(header.encode("utf-8"))

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            chunks.append("(空 sheet)\n")
            continue

        # markdown 表头：第一行作为列名
        first_row = rows[0]
        headers = [str(c) if c is not None else "" for c in first_row]
        header_line = "| " + " | ".join(headers) + " |\n"
        sep_line = "| " + " | ".join(["---"] * len(headers)) + " |\n"
        chunks.append(header_line)
        chunks.append(sep_line)
        total_bytes += len((header_line + sep_line).encode("utf-8"))

        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            line = "| " + " | ".join(cells) + " |\n"
            line_bytes = line.encode("utf-8")
            if total_bytes + len(line_bytes) > max_bytes:
                chunks.append("[已截断：剩余行/sheet 未读取]\n")
                truncated = True
                break
            chunks.append(line)
            total_bytes += len(line_bytes)

    try:
        wb.close()
    except Exception:
        pass

    return "".join(chunks).strip() or "[.xlsx 没有可读数据]"


# ──────────────── 视觉工具：渲染单页/读图为 base64 ────────────────

def _render_pdf_page_to_png_b64(abs_path: str, page_idx: int) -> tuple[Optional[str], Optional[str]]:
    """Render PDF page (0-indexed) to PNG base64. Returns (b64, error)."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return None, "[错误] pypdfium2 未安装，无法渲染 PDF"

    try:
        pdf = pdfium.PdfDocument(abs_path)
    except Exception as exc:
        return None, f"[错误] 无法打开 PDF: {exc}"

    try:
        n_pages = len(pdf)
        if page_idx < 0 or page_idx >= n_pages:
            return None, f"页码越界: PDF 共 {n_pages} 页，请求第 {page_idx + 1} 页"
        try:
            page = pdf[page_idx]
            bitmap = page.render(scale=_PDF_RENDER_SCALE)
            pil_image = bitmap.to_pil()
        except Exception as exc:
            return None, f"[错误] 渲染第 {page_idx + 1} 页失败: {exc}"

        buf = io.BytesIO()
        pil_image.save(buf, format="PNG", optimize=True)
        return base64.b64encode(buf.getvalue()).decode("ascii"), None
    finally:
        try:
            pdf.close()
        except Exception:
            pass


def _read_image_to_b64(abs_path: str) -> tuple[Optional[str], Optional[str], Optional[str]]:
    """Read image file → (base64, mime, error)."""
    try:
        size = os.path.getsize(abs_path)
    except OSError as exc:
        return None, None, f"[错误] 读取文件大小失败: {exc}"
    if size > _MAX_IMAGE_BYTES:
        return None, None, f"图片过大 ({size // 1024}KB > {_MAX_IMAGE_BYTES // 1024}KB)，请压缩后重试"

    ext = os.path.splitext(abs_path)[1].lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
    }
    mime = mime_map.get(ext, "image/png")

    try:
        with open(abs_path, "rb") as f:
            raw = f.read()
    except OSError as exc:
        return None, None, f"[错误] 读取图片失败: {exc}"
    return base64.b64encode(raw).decode("ascii"), mime, None


# ──────────────── 工具工厂 ────────────────

def create_read_document_tool(user_id: str, allowed_paths: Optional[List[str]] = None):
    """Build a ``read_document`` tool bound to the user's filesystem.

    Args:
        user_id: admin user id (consumer should pass admin_id).
        allowed_paths: optional whitelist (consumer pass ``allowed_docs``);
            None / ``["*"]`` ⇒ full filesystem.
    """
    fs_dir = get_user_filesystem_dir(user_id)

    def _is_allowed(path: str) -> bool:
        if not allowed_paths or allowed_paths == ["*"]:
            return True
        norm = path.lstrip("/").replace("\\", "/")
        for pattern in allowed_paths:
            if pattern == "*" or norm.startswith(pattern.lstrip("/")):
                return True
        return False

    @tool
    def read_document(path: str) -> str:
        """读取结构化文档（PDF / Word / Excel）并返回纯文本内容。

        **支持的格式**：
        - .pdf  → 全文文本，每页前标 [Page N]
        - .docx → 段落文本，保留标题层级（# / ## / ###）+ 表格
        - .xlsx → 每个 sheet 转 markdown 表格

        **不要用此工具读** .txt / .md / .csv / .json / .py 等纯文本格式 ——
        请用内置的 `read_file` 工具（更简单直接）。

        **超长文档自动截断在 200KB**。如果发现内容截断了：
        - PDF：用 `view_pdf_page_or_image(path, page=N)` 单页精读，或
          `run_script` 写 pypdfium2 脚本切片处理
        - .docx/.xlsx：用 `run_script` 写 python-docx / openpyxl 脚本

        Args:
            path: 文件路径（相对或绝对，会限制在你的文件系统内）
        """
        if allowed_paths and not _is_allowed(path):
            return f"无权限访问该文件: {path}"
        abs_path, err = _resolve_path(fs_dir, path)
        if err:
            return err
        assert abs_path is not None  # for type-checker

        ext = os.path.splitext(abs_path)[1].lower()

        if ext in _PLAIN_TEXT_EXTS:
            return (
                f"{path} 是纯文本格式（{ext}），请改用 `read_file` 工具，"
                "更简单且不限大小。read_document 只用于 .pdf/.docx/.xlsx 等"
                "需要解析的二进制文档。"
            )

        if ext == ".pdf":
            return _extract_pdf_text(abs_path)
        if ext == ".docx":
            return _extract_docx_text(abs_path)
        if ext == ".xlsx":
            return _extract_xlsx_text(abs_path)
        if ext == ".doc":
            return (
                "[错误] 旧版 .doc 格式（Word 97-2003）暂不支持。请将文件另存为 "
                ".docx 后重试，或用 `run_script` 调用 antiword / textract 转换。"
            )
        if ext == ".xls":
            return (
                "[错误] 旧版 .xls 格式（Excel 97-2003）暂不支持。请将文件另存为 "
                ".xlsx 后重试，或用 `run_script` + xlrd 处理。"
            )
        if ext == ".pptx":
            return "[错误] .pptx 格式暂不支持，请用 `run_script` + python-pptx 处理。"
        if ext in _IMAGE_EXTS:
            return (
                f"{path} 是图片文件，请改用 `view_pdf_page_or_image` 工具"
                "（多模态视觉读取）。"
            )

        return (
            f"[错误] 不支持的文件类型: {ext}。"
            "支持: .pdf / .docx / .xlsx；纯文本格式请用 read_file。"
        )

    return read_document


def create_view_pdf_or_image_tool(user_id: str, allowed_paths: Optional[List[str]] = None):
    """Build a ``view_pdf_page_or_image`` multimodal tool.

    Returns a list[dict] (text + image_url blocks) when invoked, which
    LangChain ToolNode wraps into a multimodal ``ToolMessage`` consumable
    by GPT-4o / Claude / Gemini etc.
    """
    fs_dir = get_user_filesystem_dir(user_id)

    def _is_allowed(path: str) -> bool:
        if not allowed_paths or allowed_paths == ["*"]:
            return True
        norm = path.lstrip("/").replace("\\", "/")
        for pattern in allowed_paths:
            if pattern == "*" or norm.startswith(pattern.lstrip("/")):
                return True
        return False

    @tool
    def view_pdf_page_or_image(path: str, page: int = 1):
        """**多模态视觉读取** — 把单页 PDF 或图片转成多模态消息让你"看到"。

        **使用场景（仅作为弥补，不要滥用）**：
        - 文本提取失败的扫描件 PDF
        - 需要查看图表 / 插图 / 公式 / 手写笔记的视觉细节
        - 单张图片的内容理解（OCR / 物体识别等）

        **❗ 重要约束**：
        - 一次只看 **一页**（PDF 必须传 page，1-based）
        - 同一文件单次对话最多调用 5 次（超出会被拒绝）
        - 想读整本 PDF 文本？请用 `read_document`，一次拿全文
        - 想精确提取大量页？请用 `run_script` 写脚本

        Args:
            path: 文件路径（PDF / png / jpg / jpeg / webp / gif / bmp）
            page: PDF 时为 1-based 页码（图片忽略此参数），默认 1
        """
        if allowed_paths and not _is_allowed(path):
            return f"无权限访问该文件: {path}"
        abs_path, err = _resolve_path(fs_dir, path)
        if err:
            return err
        assert abs_path is not None

        # throttle 检查（按 path）
        block_msg = _bump_and_check_view_count(path)
        if block_msg:
            return block_msg

        ext = os.path.splitext(abs_path)[1].lower()

        if ext in _IMAGE_EXTS:
            b64, mime, err = _read_image_to_b64(abs_path)
            if err:
                return err
            assert b64 and mime
            preface = f"图片 {path}（{mime}）的内容如下："
            return [
                {"type": "text", "text": preface},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ]

        if ext == ".pdf":
            page_idx = max(1, int(page or 1)) - 1
            b64, err = _render_pdf_page_to_png_b64(abs_path, page_idx)
            if err:
                return err
            assert b64
            preface = f"PDF {path} 第 {page_idx + 1} 页内容如下："
            return [
                {"type": "text", "text": preface},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            ]

        if ext in _PLAIN_TEXT_EXTS:
            return (
                f"{path} 是文本文件，请直接用 `read_file` 读取内容。"
                "view_pdf_page_or_image 只用于 PDF 和图片的视觉读取。"
            )
        if ext in {".docx", ".xlsx"}:
            return (
                f"{path} 是 Word/Excel 文档，请用 `read_document` 读取文本内容。"
                "view_pdf_page_or_image 只用于 PDF 单页和图片。"
            )

        return (
            f"[错误] 不支持的文件类型: {ext}。"
            "支持: .pdf（按页）/ .png / .jpg / .jpeg / .webp / .gif / .bmp"
        )

    return view_pdf_page_or_image


def create_document_tools(user_id: str, allowed_paths: Optional[List[str]] = None) -> list:
    """Convenience: return both tools as a list, ready to ``tools.extend(...)``."""
    return [
        create_read_document_tool(user_id, allowed_paths=allowed_paths),
        create_view_pdf_or_image_tool(user_id, allowed_paths=allowed_paths),
    ]
